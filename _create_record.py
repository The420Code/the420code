from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_heading('The 420 Code — Platform Registration Record', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'Last updated: {datetime.date.today().strftime("%d %B %Y")}')
doc.add_paragraph('Artist G · Studio G · Strand, Cape Town')
doc.add_paragraph('Website: https://the420code.org')
doc.add_paragraph('ORCID: 0009-0005-4873-7962')
doc.add_paragraph('')

# ===== SECTION 1 =====
doc.add_heading('1. REGISTERED & LIVE', level=1)
doc.add_paragraph('Fully set up with content uploaded and publicly accessible.')

t = doc.add_table(rows=1, cols=5)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(['Platform', 'URL', 'Status', 'Files', 'DOI / ID']):
    t.rows[0].cells[i].text = h

data1 = [
    ['the420code.org (Netlify)', 'https://the420code.org', 'LIVE', '13 pages + 200+ PDFs', 'N/A'],
    ['Zenodo — English', 'https://doi.org/10.5281/zenodo.19208226', 'PUBLISHED', '42 AP PDFs', '10.5281/zenodo.19208226'],
    ['Zenodo — Spanish', 'https://doi.org/10.5281/zenodo.19209178', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209178'],
    ['Zenodo — French', 'https://doi.org/10.5281/zenodo.19209182', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209182'],
    ['Zenodo — German', 'https://doi.org/10.5281/zenodo.19209189', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209189'],
    ['Zenodo — Portuguese', 'https://doi.org/10.5281/zenodo.19209193', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209193'],
    ['Zenodo — Dutch', 'https://doi.org/10.5281/zenodo.19209195', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209195'],
    ['Zenodo — Italian', 'https://doi.org/10.5281/zenodo.19209197', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209197'],
    ['Zenodo — Chinese', 'https://doi.org/10.5281/zenodo.19209199', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209199'],
    ['Zenodo — Japanese', 'https://doi.org/10.5281/zenodo.19209205', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209205'],
    ['Zenodo — Korean', 'https://doi.org/10.5281/zenodo.19209209', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209209'],
    ['Zenodo — Russian', 'https://doi.org/10.5281/zenodo.19209217', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209217'],
    ['Zenodo — Arabic', 'https://doi.org/10.5281/zenodo.19209221', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209221'],
    ['Zenodo — Hindi', 'https://doi.org/10.5281/zenodo.19209223', 'PUBLISHED', '5+ PDFs', '10.5281/zenodo.19209223'],
    ['Internet Archive', 'https://archive.org/details/the-420-code-artist-proofs', 'PUBLISHED', '140 PDFs', 'N/A'],
    ['OSF', 'https://osf.io/r69bm/', 'PUBLISHED', '140 PDFs', 'N/A'],
    ['ORCID', 'https://orcid.org/0009-0005-4873-7962', 'ACTIVE', 'Profile linked', '0009-0005-4873-7962'],
    ['Google Search Console', 'https://search.google.com/search-console', 'VERIFIED', '13 pages indexed', 'N/A'],
]
for r in data1:
    row = t.add_row().cells
    for i, v in enumerate(r): row[i].text = v

doc.add_paragraph('')

# ===== SECTION 2 =====
doc.add_heading('2. IN PROGRESS', level=1)
doc.add_paragraph('Accounts created, uploads started but not yet fully published.')

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Light Grid Accent 1'
for i, h in enumerate(['Platform', 'URL', 'Status', 'Notes']):
    t2.rows[0].cells[i].text = h

data2 = [
    ['figshare', 'https://figshare.com', 'UPLOADING', '80 files uploaded, processing. Check dashboard to publish.'],
    ['Google Scholar', 'https://scholar.google.com', 'PENDING', 'Citation meta tags on site. Auto-indexing via DOI.'],
    ['ResearchGate', 'https://www.researchgate.net', 'REAPPLY', 'Initial rejected. Reply sent with DOI/ORCID. Awaiting re-review.'],
]
for r in data2:
    row = t2.add_row().cells
    for i, v in enumerate(r): row[i].text = v

doc.add_paragraph('')

# ===== SECTION 3 =====
doc.add_heading('3. SEARCH ENGINES — NEED REGISTRATION', level=1)
doc.add_paragraph('Submit the site and sitemap for maximum crawler coverage.')

t3 = doc.add_table(rows=1, cols=5)
t3.style = 'Light Grid Accent 1'
for i, h in enumerate(['Search Engine', 'Region', 'Registration URL', 'Priority', 'Notes']):
    t3.rows[0].cells[i].text = h

data3 = [
    ['Bing Webmaster Tools', 'Global', 'https://www.bing.com/webmasters', 'HIGH', 'Sign in with Microsoft account. Add site + sitemap.'],
    ['Yandex Webmaster', 'Russia', 'https://webmaster.yandex.com', 'HIGH', 'Sign up. Add site. Verify with HTML tag. Submit sitemap.'],
    ['Baidu Webmaster', 'China', 'https://ziyuan.baidu.com', 'CRITICAL', 'Needs Chinese phone or partner. #1 in China.'],
    ['Sogou Webmaster', 'China', 'https://zhanzhang.sogou.com', 'HIGH', '#2 Chinese search engine.'],
    ['360/Haosou', 'China', 'https://zhanzhang.so.com', 'MEDIUM', '#3 Chinese search engine.'],
    ['Naver Search Advisor', 'Korea', 'https://searchadvisor.naver.com', 'HIGH', '#1 Korean search engine.'],
    ['Daum (Kakao)', 'Korea', 'https://register.search.daum.net', 'MEDIUM', '#2 Korean search engine.'],
]
for r in data3:
    row = t3.add_row().cells
    for i, v in enumerate(r): row[i].text = v

doc.add_paragraph('')

# ===== SECTION 4 =====
doc.add_heading('4. ACADEMIC DATABASES — NEED REGISTRATION', level=1)
doc.add_paragraph('Upload AP PDFs for maximum scholarly visibility.')

t4 = doc.add_table(rows=1, cols=5)
t4.style = 'Light Grid Accent 1'
for i, h in enumerate(['Database', 'Region', 'URL', 'Priority', 'Notes']):
    t4.rows[0].cells[i].text = h

data4 = [
    ['ChinaXiv', 'China', 'https://chinaxiv.org', 'CRITICAL', 'Chinese preprint server (CAS). Accepts independent researchers.'],
    ['CNKI', 'China', 'https://www.cnki.net', 'CRITICAL', 'Largest Chinese academic database. Needs publisher/partner.'],
    ['Wanfang Data', 'China', 'https://www.wanfangdata.com.cn', 'HIGH', '#2 Chinese academic database. Institutional.'],
    ['IndiaRxiv', 'India', 'https://indiarxiv.in', 'HIGH', 'Indian preprint server. Open to independents.'],
    ['Shodhganga', 'India', 'https://shodhganga.inflibnet.ac.in', 'MEDIUM', 'Indian repository. Institutional only.'],
    ['J-STAGE', 'Japan', 'https://www.jstage.jst.go.jp', 'HIGH', 'Japanese academic platform. Institutional.'],
    ['CiNii', 'Japan', 'https://cir.nii.ac.jp', 'HIGH', 'Japan national academic database.'],
    ['KCI', 'Korea', 'https://www.kci.go.kr', 'MEDIUM', 'Korean Citation Index. Institutional.'],
    ['CyberLeninka', 'Russia', 'https://cyberleninka.ru', 'HIGH', 'Russian open-access. Submit RU PDFs.'],
    ['eLibrary.ru', 'Russia', 'https://elibrary.ru', 'HIGH', 'Russian academic database.'],
    ['viXra', 'Global', 'https://vixra.org', 'MEDIUM', 'Physics preprints. No endorsement needed.'],
    ['PhilArchive', 'Global', 'https://philarchive.org', 'MEDIUM', 'Philosophy. Needs institutional or email request.'],
    ['SSRN', 'Global', 'https://www.ssrn.com', 'MEDIUM', 'Social sciences. Ethics/economics/policy APs.'],
    ['Academia.edu', 'Global', 'https://www.academia.edu', 'LOW', 'Researcher social network. Easy signup.'],
    ['arXiv', 'Global', 'https://arxiv.org', 'HIGH', 'Physics gold standard. Needs endorsement from existing author.'],
]
for r in data4:
    row = t4.add_row().cells
    for i, v in enumerate(r): row[i].text = v

doc.add_paragraph('')

# ===== SECTION 5 =====
doc.add_heading('5. AUTO-INDEXED (NO ACTION NEEDED)', level=1)
doc.add_paragraph('These discover content automatically from Zenodo DOIs, OSF, and Internet Archive.')

t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Light Grid Accent 1'
for i, h in enumerate(['Platform', 'Region', 'How it finds you']):
    t5.rows[0].cells[i].text = h

data5 = [
    ['Google Scholar', 'Global', 'Citation meta tags + Zenodo DOI + DataCite'],
    ['Semantic Scholar', 'Global', 'Auto-indexes from DOI via DataCite'],
    ['BASE (Bielefeld)', 'Global', 'Auto-indexes from OSF and Zenodo'],
    ['CORE (UK)', 'Global', 'Auto-indexes from Zenodo open access'],
    ['OpenAIRE', 'EU', 'Auto-indexes from Zenodo'],
    ['DataCite', 'Global', 'Zenodo mints DOIs via DataCite'],
    ['Baidu Scholar', 'China', 'Discovers via DOI and Internet Archive'],
    ['Dimensions', 'Global', 'Indexes via DOI from DataCite'],
    ['Unpaywall', 'Global', 'Finds open-access via DOI'],
    ['Yahoo Japan', 'Japan', 'Uses Google index'],
]
for r in data5:
    row = t5.add_row().cells
    for i, v in enumerate(r): row[i].text = v

doc.add_paragraph('')

# ===== SECTION 6 =====
doc.add_heading('6. SUMMARY', level=1)

t6 = doc.add_table(rows=1, cols=2)
t6.style = 'Light Grid Accent 1'
t6.rows[0].cells[0].text = 'Category'
t6.rows[0].cells[1].text = 'Count'

stats = [
    ['Platforms registered & live', '18 (incl. 13 Zenodo editions)'],
    ['Platforms in progress', '3'],
    ['Search engines to register', '7'],
    ['Academic databases to register', '15'],
    ['Auto-indexed (no action)', '10'],
    ['Total platforms (current + pending)', '53'],
]
for r in stats:
    row = t6.add_row().cells
    row[0].text = r[0]; row[1].text = r[1]

doc.add_paragraph('')
doc.add_paragraph('Document maintained by Artist G — Studio G')
doc.add_paragraph('https://the420code.org')

doc.save(r'C:\Users\info\OneDrive\00 - publish on line\420_Code_Platform_Registration_Record.docx')
print('Document saved successfully')
