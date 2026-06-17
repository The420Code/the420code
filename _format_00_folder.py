"""
Add/update last page for all DOCX files in the OneDrive 00 folder
- AP files: use existing AP metadata from _format_all_docs.py
- Non-AP books: custom metadata per book
- If last page exists: remove old text, keep logo if present, add new text + logo
"""
import sys, io, os, re, glob
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

BASE = r"C:\Users\info\OneDrive\00"
LOGO_PATH = r"C:\Users\info\OneDrive\The 420 Code - Work\Images\Eye of the Universe.jpg"

# ── AP METADATA (same as _format_all_docs.py) ───────────────
ap_titles = {
    "AP01": "The Actualization State", "AP02": "The Operator", "AP03": "The Ratio",
    "AP04": "The Loop Hypothesis", "AP05": "The Break", "AP06": "The Leakage Constant",
    "AP07": "The Record Measure", "AP08": "The Identity", "AP09": "The Break — Empty Set",
    "AP10": "The Dimension", "AP11": "The Spin", "AP12": "The Limit",
    "AP13": "The Grain", "AP14": "The Correction", "AP15": "The Connection",
    "AP16": "The Break — Electroweak", "AP17": "The Room", "AP18": "The Floor",
    "AP19": "The Direction", "AP20": "The Proof", "AP21": "The Web",
    "AP22": "The Ledger", "AP23": "The Single Record", "AP24": "The Residual",
    "AP25": "The Measure", "AP26": "The Surplus", "AP27": "The Harmonics",
    "AP28": "The Constant", "AP29": "The Actualization Proof", "AP30": "The Resistance",
    "AP31": "The Alignment", "AP32": "The Correction — Ethics", "AP33": "The Boundary",
    "AP34": "The Inversion", "AP35": "The Ledger — Economics", "AP36": "The Feed",
    "AP37": "The First Boundary", "AP38": "The Exit", "AP39": "The Scaffold",
    "AP40": "The Irrational", "AP41": "The Loop", "AP42": "The Clock",
}

ap_notebooks = {
    "AP01": "Ø.1 The Premise", "AP02": "Ø.7 Consciousness & The Ethic",
    "AP03": "Ø.2 Spacetime", "AP04": "Ø.6 Cosmology", "AP05": "Ø.2 Spacetime",
    "AP06": "Ø.4 Forces & Constants", "AP07": "Ø.3 Quantum Mechanics",
    "AP08": "Ø.2 Spacetime", "AP09": "Ø.3 Quantum Mechanics",
    "AP10": "Ø.2 Spacetime", "AP11": "Ø.3 Quantum Mechanics",
    "AP12": "Ø.3 Quantum Mechanics", "AP13": "Ø.3 Quantum Mechanics",
    "AP14": "Ø.4 Forces & Constants", "AP15": "Ø.4 Forces & Constants",
    "AP16": "Ø.4 Forces & Constants", "AP17": "Ø.6 Cosmology",
    "AP18": "Ø.6 Cosmology", "AP19": "Ø.4 Forces & Constants",
    "AP20": "Ø.1 The Premise", "AP21": "Ø.6 Cosmology",
    "AP22": "Ø.5 Particles & Matter", "AP23": "Ø.3 Quantum Mechanics",
    "AP24": "Ø.4 Forces & Constants", "AP25": "Ø.3 Quantum Mechanics",
    "AP26": "Ø.5 Particles & Matter", "AP27": "Ø.4 Forces & Constants",
    "AP28": "Ø.5 Particles & Matter", "AP29": "Ø.7 Consciousness & The Ethic",
    "AP30": "Ø.5 Particles & Matter", "AP31": "Ø.7 Consciousness & The Ethic",
    "AP32": "Ø.7 Consciousness & The Ethic", "AP33": "Ø.7 Consciousness & The Ethic",
    "AP34": "Ø.8 Applications", "AP35": "Ø.8 Applications",
    "AP36": "Ø.8 Applications", "AP37": "Ø.8 Applications",
    "AP38": "Ø.7 Consciousness & The Ethic", "AP39": "Ø.7 Consciousness & The Ethic",
    "AP40": "Ø.1 The Premise", "AP41": "Ø.6 Cosmology", "AP42": "Ø.6 Cosmology",
}

# ── NON-AP BOOK METADATA ────────────────────────────────────
# Maps filename patterns to (series_value, volume, edition, title)
book_metadata = {
    "Illusion of the Other": ("The 420 Code", "The Exhibition", "Edition 01", "The Illusion of the Other"),
    "Rosin Ø Prose": ("The 420 Code", "The Rosin", "Edition 01 — Prose", "The Rosin — Prose"),
    "Rosin Ø Proofs": ("The 420 Code", "The Rosin", "Edition 05 — Proofs", "The Rosin — Proofs"),
    "The_Rosin_Formatted": ("The 420 Code", "The Rosin", "Edition 01", "The Rosin"),
    "Editions Ø Prose": ("The 420 Code", "The Editions", "Edition 01 — Prose", "The Editions — Prose"),
    "Notebooks Ø Prose": ("The 420 Code", "The Notebooks", "Edition 01 — Prose", "The Notebooks — Prose"),
    "Notebook Ø.1 The Premise": ("The 420 Code", "Notebook Ø.1", "Edition 01", "The Premise"),
    "Notebook Ø.2 Spacetime": ("The 420 Code", "Notebook Ø.2", "Edition 01", "Spacetime"),
    "Notebook Ø.3 Quantum Mechanics": ("The 420 Code", "Notebook Ø.3", "Edition 01", "Quantum Mechanics"),
    "Notebook Ø.4 Forces and Constants": ("The 420 Code", "Notebook Ø.4", "Edition 01", "Forces and Constants"),
    "Notebook Ø.5 Particles_and_Matter": ("The 420 Code", "Notebook Ø.5", "Edition 01", "Particles and Matter"),
    "Notebook Ø.6 Cosmology": ("The 420 Code", "Notebook Ø.6", "Edition 01", "Cosmology"),
    "Notebook Ø.7 Consciousness and the Ethic": ("The 420 Code", "Notebook Ø.7", "Edition 01", "Consciousness and the Ethic"),
    "Notebook Ø.8 Applications": ("The 420 Code", "Notebook Ø.8", "Edition 01", "Applications"),
    "Are_You_Certain_FINAL_PROOFED": ("The 420 Code", "The Records", "Record 03", "Are You Certain"),
    "Are_You_Certain - final to format": ("The 420 Code", "The Records", "Record 03", "Are You Certain"),
    "Master Kill Switch Registry - 2603": ("The 420 Code", "Reference", "Edition 01", "Master Kill Switch Registry"),
    "Master_Kill_Switch_Registry_v4(1)": ("The 420 Code", "Reference", "Edition 01", "Master Kill Switch Registry"),
    "Artist Proofs all of them": ("The 420 Code", "The Notebooks", "Complete", "Artist Proofs — Complete Collection"),
    "The Notebooks – AP01-42": ("The 420 Code", "The Notebooks", "Complete", "The Notebooks — AP01–42"),
}

LAST_PAGE_LINES = [
    "This work is published for free forever at www.the420code.org.",
    "No paywall. No institutional affiliation. No gatekeepers.",
    "This work is Copyleft. You are free to download, print, share, and distribute. You are not free to alter the source. Keep the signal clean.",
]


def find_last_page_index(doc):
    """Find the paragraph index of 'Last Page' heading."""
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        if doc.paragraphs[i].text.strip() == "Last Page":
            return i
    return -1


def remove_from_index(doc, start_idx):
    """Remove all paragraphs from start_idx to end."""
    # Go back one if previous paragraph is blank (page break)
    if start_idx > 0 and not doc.paragraphs[start_idx - 1].text.strip():
        # Check if it has a page break
        start_idx -= 1

    body = doc.element.body
    elems = [doc.paragraphs[i]._element for i in range(start_idx, len(doc.paragraphs))]
    for elem in elems:
        body.remove(elem)


def add_last_page_ap(doc, ap_num):
    """Add last page for an AP document."""
    title = ap_titles.get(ap_num, "Unknown")
    notebook = ap_notebooks.get(ap_num, "Unknown")
    ap_int = int(ap_num[2:])

    _add_last_page_content(doc, [
        ("Series", "The 420 Code"),
        ("Volume", f"Notebook {notebook}"),
        ("Edition", f"Artist Proof {ap_int:02d}"),
        ("Title", title),
        ("Artist", "G"),
        ("Date", "March 2026"),
    ])


def add_last_page_book(doc, series, volume, edition, title):
    """Add last page for a non-AP book."""
    _add_last_page_content(doc, [
        ("Series", series),
        ("Volume", volume),
        ("Edition", edition),
        ("Title", title),
        ("Artist", "G"),
        ("Date", "March 2026"),
    ])


def _add_last_page_content(doc, meta_lines):
    """Shared last page builder."""
    # Page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # Heading
    p = doc.add_paragraph()
    run = p.add_run("Last Page")
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()

    # Metadata
    for label, value in meta_lines:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}\t{value}")
        run.font.bold = True

    doc.add_paragraph()

    # Prose
    for line in LAST_PAGE_LINES:
        p = doc.add_paragraph()
        p.add_run(line)

    doc.add_paragraph()

    # Logo
    p = doc.add_paragraph()
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(3.2), height=Cm(0.8))
    except Exception:
        pass

    doc.add_paragraph()


def process_file(filepath):
    """Process a single file."""
    fname = os.path.basename(filepath)
    fname_no_ext = os.path.splitext(fname)[0]

    try:
        doc = Document(filepath)
    except Exception as e:
        return f"ERROR: {fname[:60]}: {e}"

    # Detect AP number
    m = re.search(r'AP(\d+)', fname)
    ap_num = f"AP{int(m.group(1)):02d}" if m else None

    # Remove existing last page if present
    lp_idx = find_last_page_index(doc)
    if lp_idx >= 0:
        remove_from_index(doc, lp_idx)

    # Add new last page
    if ap_num and ap_num in ap_titles:
        add_last_page_ap(doc, ap_num)
        tag = f"AP:{ap_num}"
    else:
        # Match to book metadata
        matched = False
        for pattern, (series, volume, edition, title) in book_metadata.items():
            if pattern in fname_no_ext:
                add_last_page_book(doc, series, volume, edition, title)
                tag = f"BOOK:{title[:30]}"
                matched = True
                break

        if not matched:
            # Generic fallback
            add_last_page_book(doc, "The 420 Code", "—", "—", fname_no_ext[:50])
            tag = f"GENERIC:{fname_no_ext[:30]}"

    doc.save(filepath)
    return f"OK [{tag}] {fname[:70]}"


# ── MAIN ─────────────────────────────────────────────────────
if __name__ == "__main__":
    files = sorted(glob.glob(os.path.join(BASE, "*.docx")))
    print(f"Processing {len(files)} files in {BASE}\n")

    ok = err = 0
    for f in files:
        result = process_file(f)
        print(f"  {result}")
        if result.startswith("OK"):
            ok += 1
        else:
            err += 1

    print(f"\n=== DONE === OK: {ok} | Errors: {err}")
