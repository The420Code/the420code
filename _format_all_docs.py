"""
Standardize formatting and insert translated last page for all AP documents.
Template: AP01 master DOCX formatting.
"""
import sys, os, io, re, glob, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"C:\Users\info\OneDrive\00 - publish on line"

# ── AP METADATA ──────────────────────────────────────────────
ap_titles = {
    "AP01": "The Actualization State", "AP02": "The Operator", "AP03": "The Ratio",
    "AP04": "The Loop Hypothesis", "AP05": "The Break", "AP06": "The Leakage Constant",
    "AP07": "The Record Measure", "AP08": "The Identity", "AP09": "The Break — Empty Set",
    "AP10": "The Dimension", "AP11": "The Spin", "AP12": "The Limit",
    "AP13": "The Grain", "AP14": "The Correction", "AP15": "The Connection",
    "AP16": "The Break — Electroweak", "AP17": "The Room", "AP18": "The Floor",
    "AP19": "The Direction", "AP20": "The Proof", "AP21": "The Web",
    "AP22": "The Ledger", "AP23": "The Single Record", "AP24": "The Residual",
    "AP25": "The Measure", "AP26": "The Surplus", "AP27": "The Harmonics",
    "AP28": "The Constant", "AP29": "The Actualization Proof", "AP30": "The Resistance",
    "AP31": "The Alignment", "AP32": "The Correction — Ethics", "AP33": "The Boundary",
    "AP34": "The Inversion", "AP35": "The Ledger — Economics", "AP36": "The Feed",
    "AP37": "The First Boundary", "AP38": "The Exit", "AP39": "The Scaffold",
    "AP40": "The Irrational", "AP41": "The Loop", "AP42": "The Clock",
}

ap_notebooks = {
    "AP01": "Ø.1 The Premise", "AP02": "Ø.7 Consciousness & The Ethic",
    "AP03": "Ø.2 Spacetime", "AP04": "Ø.6 Cosmology", "AP05": "Ø.2 Spacetime",
    "AP06": "Ø.4 Forces & Constants", "AP07": "Ø.3 Quantum Mechanics",
    "AP08": "Ø.2 Spacetime", "AP09": "Ø.3 Quantum Mechanics",
    "AP10": "Ø.2 Spacetime", "AP11": "Ø.3 Quantum Mechanics",
    "AP12": "Ø.3 Quantum Mechanics", "AP13": "Ø.3 Quantum Mechanics",
    "AP14": "Ø.4 Forces & Constants", "AP15": "Ø.4 Forces & Constants",
    "AP16": "Ø.4 Forces & Constants", "AP17": "Ø.6 Cosmology",
    "AP18": "Ø.6 Cosmology", "AP19": "Ø.4 Forces & Constants",
    "AP20": "Ø.1 The Premise", "AP21": "Ø.6 Cosmology",
    "AP22": "Ø.5 Particles & Matter", "AP23": "Ø.3 Quantum Mechanics",
    "AP24": "Ø.4 Forces & Constants", "AP25": "Ø.3 Quantum Mechanics",
    "AP26": "Ø.5 Particles & Matter", "AP27": "Ø.4 Forces & Constants",
    "AP28": "Ø.5 Particles & Matter", "AP29": "Ø.7 Consciousness & The Ethic",
    "AP30": "Ø.5 Particles & Matter", "AP31": "Ø.7 Consciousness & The Ethic",
    "AP32": "Ø.7 Consciousness & The Ethic", "AP33": "Ø.7 Consciousness & The Ethic",
    "AP34": "Ø.8 Applications", "AP35": "Ø.8 Applications",
    "AP36": "Ø.8 Applications", "AP37": "Ø.8 Applications",
    "AP38": "Ø.7 Consciousness & The Ethic", "AP39": "Ø.7 Consciousness & The Ethic",
    "AP40": "Ø.1 The Premise", "AP41": "Ø.6 Cosmology", "AP42": "Ø.6 Cosmology",
}

# ── LAST PAGE TRANSLATIONS ───────────────────────────────────
# Each language gets the full last page text block
last_page_text = {
    "EN": {
        "heading": "Last Page",
        "series_label": "Series", "series_value": "The 420 Code",
        "volume_label": "Volume", "edition_label": "Edition", "title_label": "Title",
        "artist_label": "Artist", "artist_value": "G",
        "date_label": "Date", "date_value": "March 2026",
        "lines": [
            "This work is published for free forever at www.the420code.org.",
            "No paywall. No institutional affiliation. No gatekeepers.",
            "This work is Copyleft. You are free to download, print, share, and distribute. You are not free to alter the source. Keep the signal clean.",
        ],
    },
    "DE": {
        "heading": "Letzte Seite",
        "series_label": "Reihe", "series_value": "The 420 Code",
        "volume_label": "Band", "edition_label": "Ausgabe", "title_label": "Titel",
        "artist_label": "Künstler", "artist_value": "G",
        "date_label": "Datum", "date_value": "März 2026",
        "lines": [
            "Dieses Werk ist für immer kostenlos veröffentlicht unter www.the420code.org.",
            "Keine Paywall. Keine institutionelle Zugehörigkeit. Keine Torwächter.",
            "Dieses Werk ist Copyleft. Sie dürfen es herunterladen, drucken, teilen und verbreiten. Sie dürfen die Quelle nicht verändern. Halten Sie das Signal sauber.",
        ],
    },
    "ES": {
        "heading": "Última Página",
        "series_label": "Serie", "series_value": "The 420 Code",
        "volume_label": "Volumen", "edition_label": "Edición", "title_label": "Título",
        "artist_label": "Artista", "artist_value": "G",
        "date_label": "Fecha", "date_value": "Marzo 2026",
        "lines": [
            "Esta obra se publica de forma gratuita para siempre en www.the420code.org.",
            "Sin muro de pago. Sin afiliación institucional. Sin guardianes.",
            "Esta obra es Copyleft. Eres libre de descargar, imprimir, compartir y distribuir. No eres libre de alterar la fuente. Mantén la señal limpia.",
        ],
    },
    "FR": {
        "heading": "Dernière Page",
        "series_label": "Série", "series_value": "The 420 Code",
        "volume_label": "Volume", "edition_label": "Édition", "title_label": "Titre",
        "artist_label": "Artiste", "artist_value": "G",
        "date_label": "Date", "date_value": "Mars 2026",
        "lines": [
            "Cette œuvre est publiée gratuitement pour toujours sur www.the420code.org.",
            "Pas de paywall. Pas d'affiliation institutionnelle. Pas de gardiens.",
            "Cette œuvre est Copyleft. Vous êtes libre de télécharger, imprimer, partager et distribuer. Vous n'êtes pas libre de modifier la source. Gardez le signal propre.",
        ],
    },
    "PT": {
        "heading": "Última Página",
        "series_label": "Série", "series_value": "The 420 Code",
        "volume_label": "Volume", "edition_label": "Edição", "title_label": "Título",
        "artist_label": "Artista", "artist_value": "G",
        "date_label": "Data", "date_value": "Março 2026",
        "lines": [
            "Esta obra é publicada gratuitamente para sempre em www.the420code.org.",
            "Sem paywall. Sem afiliação institucional. Sem porteiros.",
            "Esta obra é Copyleft. Você é livre para baixar, imprimir, compartilhar e distribuir. Você não é livre para alterar a fonte. Mantenha o sinal limpo.",
        ],
    },
    "NL": {
        "heading": "Laatste Pagina",
        "series_label": "Reeks", "series_value": "The 420 Code",
        "volume_label": "Deel", "edition_label": "Editie", "title_label": "Titel",
        "artist_label": "Kunstenaar", "artist_value": "G",
        "date_label": "Datum", "date_value": "Maart 2026",
        "lines": [
            "Dit werk is voor altijd gratis gepubliceerd op www.the420code.org.",
            "Geen betaalmuur. Geen institutionele affiliatie. Geen poortwachters.",
            "Dit werk is Copyleft. U bent vrij om te downloaden, af te drukken, te delen en te verspreiden. U bent niet vrij om de bron te wijzigen. Houd het signaal schoon.",
        ],
    },
    "IT": {
        "heading": "Ultima Pagina",
        "series_label": "Serie", "series_value": "The 420 Code",
        "volume_label": "Volume", "edition_label": "Edizione", "title_label": "Titolo",
        "artist_label": "Artista", "artist_value": "G",
        "date_label": "Data", "date_value": "Marzo 2026",
        "lines": [
            "Quest'opera è pubblicata gratuitamente per sempre su www.the420code.org.",
            "Nessun paywall. Nessuna affiliazione istituzionale. Nessun guardiano.",
            "Quest'opera è Copyleft. Sei libero di scaricare, stampare, condividere e distribuire. Non sei libero di alterare la fonte. Mantieni il segnale pulito.",
        ],
    },
    "ZH": {
        "heading": "最后一页",
        "series_label": "系列", "series_value": "The 420 Code",
        "volume_label": "卷", "edition_label": "版本", "title_label": "标题",
        "artist_label": "作者", "artist_value": "G",
        "date_label": "日期", "date_value": "2026年3月",
        "lines": [
            "本作品永久免费发布于 www.the420code.org。",
            "没有付费墙。没有机构隶属。没有看门人。",
            "本作品为Copyleft。你可以自由下载、打印、分享和分发。你不可以修改源文件。保持信号干净。",
        ],
    },
    "JA": {
        "heading": "最終ページ",
        "series_label": "シリーズ", "series_value": "The 420 Code",
        "volume_label": "巻", "edition_label": "版", "title_label": "タイトル",
        "artist_label": "アーティスト", "artist_value": "G",
        "date_label": "日付", "date_value": "2026年3月",
        "lines": [
            "この作品は www.the420code.org にて永久に無料で公開されています。",
            "ペイウォールなし。機関所属なし。門番なし。",
            "この作品はコピーレフトです。ダウンロード、印刷、共有、配布は自由です。ソースの改変は許可されていません。信号をクリーンに保ってください。",
        ],
    },
    "KO": {
        "heading": "마지막 페이지",
        "series_label": "시리즈", "series_value": "The 420 Code",
        "volume_label": "권", "edition_label": "판", "title_label": "제목",
        "artist_label": "아티스트", "artist_value": "G",
        "date_label": "날짜", "date_value": "2026년 3월",
        "lines": [
            "이 작품은 www.the420code.org에서 영원히 무료로 공개됩니다.",
            "유료 장벽 없음. 기관 소속 없음. 문지기 없음.",
            "이 작품은 카피레프트입니다. 다운로드, 인쇄, 공유, 배포가 자유롭습니다. 소스를 변경할 수는 없습니다. 신호를 깨끗하게 유지하세요.",
        ],
    },
    "RU": {
        "heading": "Последняя страница",
        "series_label": "Серия", "series_value": "The 420 Code",
        "volume_label": "Том", "edition_label": "Издание", "title_label": "Название",
        "artist_label": "Автор", "artist_value": "G",
        "date_label": "Дата", "date_value": "Март 2026",
        "lines": [
            "Эта работа опубликована бесплатно навсегда на www.the420code.org.",
            "Без платного доступа. Без институциональной принадлежности. Без привратников.",
            "Эта работа — Copyleft. Вы можете свободно скачивать, печатать, делиться и распространять. Вы не можете изменять источник. Сохраняйте сигнал чистым.",
        ],
    },
    "AR": {
        "heading": "الصفحة الأخيرة",
        "series_label": "السلسلة", "series_value": "The 420 Code",
        "volume_label": "المجلد", "edition_label": "الإصدار", "title_label": "العنوان",
        "artist_label": "الفنان", "artist_value": "G",
        "date_label": "التاريخ", "date_value": "مارس ٢٠٢٦",
        "lines": [
            "هذا العمل منشور مجانًا إلى الأبد على www.the420code.org.",
            "لا جدار دفع. لا انتماء مؤسسي. لا حراس بوابات.",
            "هذا العمل هو Copyleft. أنت حر في التحميل والطباعة والمشاركة والتوزيع. لست حرًا في تغيير المصدر. حافظ على الإشارة نظيفة.",
        ],
    },
    "HI": {
        "heading": "अंतिम पृष्ठ",
        "series_label": "श्रृंखला", "series_value": "The 420 Code",
        "volume_label": "खंड", "edition_label": "संस्करण", "title_label": "शीर्षक",
        "artist_label": "कलाकार", "artist_value": "G",
        "date_label": "तिथि", "date_value": "मार्च 2026",
        "lines": [
            "यह कार्य www.the420code.org पर हमेशा के लिए मुफ्त प्रकाशित है।",
            "कोई पेवॉल नहीं। कोई संस्थागत संबद्धता नहीं। कोई द्वारपाल नहीं।",
            "यह कार्य कॉपीलेफ्ट है। आप डाउनलोड, प्रिंट, साझा और वितरित करने के लिए स्वतंत्र हैं। आप स्रोत को बदलने के लिए स्वतंत्र नहीं हैं। संकेत को स्वच्छ रखें।",
        ],
    },
}


def detect_language(filename):
    """Detect language from filename."""
    fn = os.path.basename(filename)
    fn_upper = fn.upper()
    # Check 2-letter codes at end of filename (before extension) — most reliable
    # Match patterns like _NL.docx, _DE.docx, _AR.docx
    m = re.search(r'_([A-Z]{2})\.\w+$', fn)
    if m:
        code = m.group(1)
        if code in ["AR", "DE", "ES", "FR", "HI", "IT", "JA", "KO", "NL", "PT", "RU", "ZH", "EN"]:
            return code
    # Check for full language names
    name_map = {"ARABIC": "AR", "CHINESE": "ZH", "MANDARIN": "ZH", "DUTCH": "NL",
                "FRENCH": "FR", "GERMAN": "DE", "HINDI": "HI", "ITALIAN": "IT",
                "JAPANESE": "JA", "KOREAN": "KO", "PORTUGUESE": "PT", "RUSSIAN": "RU",
                "SPANISH": "ES"}
    for name, code in name_map.items():
        if name in fn_upper:
            return code
    # Fallback: check for _XX_ in the middle
    for code in ["AR", "DE", "ES", "FR", "HI", "IT", "JA", "KO", "NL", "PT", "RU", "ZH"]:
        if f"_{code}_" in fn_upper:
            return code
    return "EN"


def detect_ap(filename):
    """Detect AP number from filename."""
    m = re.search(r'AP(\d+)', os.path.basename(filename))
    if m:
        return f"AP{int(m.group(1)):02d}"
    return None


LOGO_PATH = r"C:\Users\info\OneDrive\The 420 Code - Work\Images\Eye of the Universe.jpg"

def add_last_page(doc, ap_num, lang):
    """Add the translated last page to the document."""
    from docx.enum.text import WD_BREAK

    t = last_page_text.get(lang, last_page_text["EN"])
    title = ap_titles.get(ap_num, "Unknown")
    notebook = ap_notebooks.get(ap_num, "Unknown")
    ap_int = int(ap_num[2:])

    # Page break before last page
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # "Last Page" heading
    p = doc.add_paragraph()
    run = p.add_run(t["heading"])
    run.font.size = Pt(20)
    run.font.bold = True

    doc.add_paragraph()  # blank line

    # Metadata table (using tab-separated lines)
    meta_lines = [
        (t["series_label"], t["series_value"]),
        (t["volume_label"], f"Notebook {notebook}"),
        (t["edition_label"], f"Artist Proof {ap_int:02d}"),
        (t["title_label"], title),
        (t["artist_label"], t["artist_value"]),
        (t["date_label"], t["date_value"]),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}\t{value}")
        run.font.bold = True

    doc.add_paragraph()  # blank line

    # Three prose lines (simplified template)
    for line in t["lines"]:
        p = doc.add_paragraph()
        run = p.add_run(line)

    doc.add_paragraph()  # blank line

    # Studio G logo image (3.2cm x 0.8cm)
    p = doc.add_paragraph()
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(3.2), height=Cm(0.8))
    except Exception:
        pass  # skip image if file not found

    doc.add_paragraph()  # final blank


def has_last_page(doc):
    """Check if document already has a 'Last Page' section."""
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in [t["heading"] for t in last_page_text.values()]:
            return i
    # Also check for the English original
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Last Page":
            return i
    return -1


def remove_last_page(doc, start_idx):
    """Remove everything from last page heading onwards."""
    # We need to remove paragraphs from start_idx to end
    # But also remove a preceding page break paragraph if present
    if start_idx > 0 and not doc.paragraphs[start_idx - 1].text.strip():
        start_idx -= 1

    body = doc.element.body
    paras_to_remove = []
    for i in range(start_idx, len(doc.paragraphs)):
        paras_to_remove.append(doc.paragraphs[i]._element)

    for elem in paras_to_remove:
        body.remove(elem)


def process_file(filepath, dry_run=False):
    """Process a single DOCX file."""
    ap_num = detect_ap(filepath)
    lang = detect_language(filepath)

    if not ap_num:
        return f"SKIP (no AP number): {os.path.basename(filepath)}"

    if ap_num not in ap_titles:
        return f"SKIP (unknown AP): {os.path.basename(filepath)}"

    try:
        doc = Document(filepath)
    except Exception as e:
        return f"ERROR opening {os.path.basename(filepath)}: {e}"

    # Check/remove existing last page
    lp_idx = has_last_page(doc)
    if lp_idx >= 0:
        remove_last_page(doc, lp_idx)

    # Add new last page
    add_last_page(doc, ap_num, lang)

    if not dry_run:
        doc.save(filepath)
        return f"OK [{lang}] {os.path.basename(filepath)}"
    else:
        return f"DRY [{lang}] {os.path.basename(filepath)}"


# ── MAIN ─────────────────────────────────────────────────────
if __name__ == "__main__":
    dry_run = "--dry-run" in sys.argv

    if dry_run:
        print("=== DRY RUN — no files will be modified ===\n")

    results = {"ok": 0, "skip": 0, "error": 0}

    # Process all DOCX files in AP folders
    for i in range(1, 43):
        ap = f"AP{i:02d}"
        # Find the folder
        pattern = os.path.join(BASE, f"{ap} -*")
        folders = glob.glob(pattern)
        if not folders:
            print(f"  SKIP: No folder for {ap}")
            results["skip"] += 1
            continue

        folder = folders[0]
        docx_files = glob.glob(os.path.join(folder, "*.docx"))

        if not docx_files:
            print(f"  SKIP: No DOCX in {os.path.basename(folder)}")
            results["skip"] += 1
            continue

        for f in sorted(docx_files):
            result = process_file(f, dry_run)
            print(f"  {result}")
            if result.startswith("OK") or result.startswith("DRY"):
                results["ok"] += 1
            elif result.startswith("SKIP"):
                results["skip"] += 1
            else:
                results["error"] += 1

    print(f"\n=== DONE === OK: {results['ok']} | Skipped: {results['skip']} | Errors: {results['error']}")
